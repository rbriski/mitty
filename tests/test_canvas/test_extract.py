"""Tests for mitty.canvas.extract — file download + text extraction."""

from __future__ import annotations

import io
from unittest.mock import AsyncMock, MagicMock

import httpx
import pytest

from mitty.canvas.extract import (
    ALLOWED_CANVAS_HOSTS,
    MAX_FILE_SIZE,
    download_file_content,
    extract_text,
    extract_text_from_docx,
    extract_text_from_pdf,
    pdf_pages_to_images,
    validate_canvas_url,
)

# ---------------------------------------------------------------------------
# Helpers: create minimal PDF and DOCX in memory
# ---------------------------------------------------------------------------


def _make_pdf(text: str = "Hello from PDF") -> bytes:
    """Create a minimal single-page PDF containing *text*."""
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def _make_docx(text: str = "Hello from DOCX") -> bytes:
    """Create a minimal DOCX containing a single paragraph with *text*."""
    import docx

    doc = docx.Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------


class TestExtractPdf:
    """extract_text_from_pdf uses pymupdf to read PDF content."""

    def test_extract_pdf_returns_text(self) -> None:
        pdf_bytes = _make_pdf("Chapter 1: Introduction")
        result = extract_text_from_pdf(pdf_bytes)
        assert "Chapter 1: Introduction" in result

    def test_extract_pdf_multi_page(self) -> None:
        """Text from all pages is included."""
        import pymupdf

        doc = pymupdf.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i + 1} content")
        pdf_bytes = doc.tobytes()
        doc.close()

        result = extract_text_from_pdf(pdf_bytes)
        assert "Page 1 content" in result
        assert "Page 2 content" in result
        assert "Page 3 content" in result

    def test_extract_empty_bytes_returns_empty_string(self) -> None:
        assert extract_text_from_pdf(b"") == ""

    def test_extract_corrupt_bytes_returns_empty_string(self) -> None:
        assert extract_text_from_pdf(b"this is not a pdf") == ""


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------


class TestExtractDocx:
    """extract_text_from_docx uses python-docx to read Word content."""

    def test_extract_docx_returns_text(self) -> None:
        docx_bytes = _make_docx("Study Guide for Midterm")
        result = extract_text_from_docx(docx_bytes)
        assert "Study Guide for Midterm" in result

    def test_extract_docx_multi_paragraph(self) -> None:
        """All non-empty paragraphs are included."""
        import docx

        doc = docx.Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("")  # empty — should be skipped
        doc.add_paragraph("Third paragraph")
        buf = io.BytesIO()
        doc.save(buf)

        result = extract_text_from_docx(buf.getvalue())
        assert "First paragraph" in result
        assert "Third paragraph" in result

    def test_extract_empty_bytes_returns_empty_string(self) -> None:
        assert extract_text_from_docx(b"") == ""

    def test_extract_corrupt_bytes_returns_empty_string(self) -> None:
        assert extract_text_from_docx(b"not a docx file") == ""


# ---------------------------------------------------------------------------
# extract_text dispatcher
# ---------------------------------------------------------------------------


class TestExtractText:
    """extract_text dispatches to the correct extractor by content type."""

    def test_dispatches_pdf(self) -> None:
        pdf_bytes = _make_pdf("PDF dispatch test")
        result = extract_text(pdf_bytes, "application/pdf")
        assert "PDF dispatch test" in result

    def test_dispatches_docx(self) -> None:
        docx_bytes = _make_docx("DOCX dispatch test")
        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        result = extract_text(docx_bytes, content_type)
        assert "DOCX dispatch test" in result

    def test_unsupported_type_returns_empty(self) -> None:
        assert extract_text(b"data", "image/png") == ""

    def test_empty_content_type_returns_empty(self) -> None:
        assert extract_text(b"data", "") == ""


# ---------------------------------------------------------------------------
# URL validation
# ---------------------------------------------------------------------------


class TestValidateCanvasUrl:
    """validate_canvas_url checks hostnames against the allow list."""

    @pytest.mark.parametrize(
        "url",
        [
            "https://mitty.instructure.com/files/123/download",
            "https://canvas.instructure.com/files/456/download",
            "https://instructure-uploads.s3.amazonaws.com/bucket/key",
        ],
    )
    def test_validate_canvas_url_accepts_valid(self, url: str) -> None:
        assert validate_canvas_url(url) is True

    @pytest.mark.parametrize(
        "url",
        [
            "https://evil.com/files/123/download",
            "https://example.com/malicious",
            "ftp://mitty.instructure.com/files/1",
            "https://not-canvas.instructure.com/files/1",
            "",
            "not-a-url",
        ],
    )
    def test_validate_canvas_url_rejects_invalid(self, url: str) -> None:
        assert validate_canvas_url(url) is False

    def test_allowed_hosts_constant_is_set(self) -> None:
        """Ensure the constant contains expected hosts."""
        assert "mitty.instructure.com" in ALLOWED_CANVAS_HOSTS
        assert "canvas.instructure.com" in ALLOWED_CANVAS_HOSTS
        assert "instructure-uploads.s3.amazonaws.com" in ALLOWED_CANVAS_HOSTS


# ---------------------------------------------------------------------------
# File download
# ---------------------------------------------------------------------------


class TestDownloadFileContent:
    """download_file_content validates URLs and enforces size limits."""

    async def test_download_validates_hostname(self) -> None:
        """Rejects download from disallowed hostname."""
        client = AsyncMock(spec=httpx.AsyncClient)
        result = await download_file_content(client, "https://evil.com/malware.pdf")
        assert result is None
        client.get.assert_not_called()

    async def test_download_success(self) -> None:
        """Successful download returns content bytes."""
        content = b"fake pdf content"
        mock_response = MagicMock(spec=httpx.Response)
        mock_response.content = content
        mock_response.status_code = 200
        mock_response.raise_for_status = MagicMock()

        client = AsyncMock(spec=httpx.AsyncClient)
        client.get = AsyncMock(return_value=mock_response)

        result = await download_file_content(
            client,
            "https://mitty.instructure.com/files/123/download",
        )
        assert result == content
        client.get.assert_called_once()

    async def test_download_respects_size_limit(self) -> None:
        """Files exceeding max_size are rejected."""
        oversized = b"x" * 200  # 200 bytes
        mock_response = MagicMock(spec=httpx.Response)
        mock_response.content = oversized
        mock_response.status_code = 200
        mock_response.raise_for_status = MagicMock()

        client = AsyncMock(spec=httpx.AsyncClient)
        client.get = AsyncMock(return_value=mock_response)

        result = await download_file_content(
            client,
            "https://mitty.instructure.com/files/1/download",
            max_size=100,  # smaller than content
        )
        assert result is None

    async def test_download_http_error_returns_none(self) -> None:
        """HTTP errors are caught and return None."""
        client = AsyncMock(spec=httpx.AsyncClient)
        client.get = AsyncMock(
            side_effect=httpx.HTTPStatusError(
                "404",
                request=MagicMock(),
                response=MagicMock(),
            )
        )

        result = await download_file_content(
            client,
            "https://mitty.instructure.com/files/999/download",
        )
        assert result is None

    async def test_download_timeout_returns_none(self) -> None:
        """Timeouts are caught and return None."""
        client = AsyncMock(spec=httpx.AsyncClient)
        client.get = AsyncMock(side_effect=httpx.ReadTimeout("timed out"))

        result = await download_file_content(
            client,
            "https://mitty.instructure.com/files/1/download",
        )
        assert result is None

    async def test_max_file_size_constant(self) -> None:
        """MAX_FILE_SIZE is 10 MB."""
        assert MAX_FILE_SIZE == 10 * 1024 * 1024


# ---------------------------------------------------------------------------
# Helpers: multi-page PDF builder
# ---------------------------------------------------------------------------


def _make_pdf_pages(num_pages: int, text_prefix: str = "Page") -> bytes:
    """Create a PDF with *num_pages* pages, each containing text."""
    import pymupdf

    doc = pymupdf.open()
    for i in range(num_pages):
        page = doc.new_page()
        page.insert_text((72, 72), f"{text_prefix} {i + 1}")
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


# ---------------------------------------------------------------------------
# PDF-to-images extraction
# ---------------------------------------------------------------------------


class TestPdfPagesToImages:
    """pdf_pages_to_images converts PDF pages to PNG bytes."""

    def test_single_page(self) -> None:
        """A single-page PDF produces exactly 1 PNG."""
        pdf_bytes = _make_pdf("Single page content")
        result = pdf_pages_to_images(pdf_bytes)
        assert len(result) == 1
        # Verify it's a valid PNG (starts with PNG magic bytes)
        assert result[0][:8] == b"\x89PNG\r\n\x1a\n"

    def test_multi_page(self) -> None:
        """A multi-page PDF returns the correct number of images."""
        pdf_bytes = _make_pdf_pages(5)
        result = pdf_pages_to_images(pdf_bytes)
        assert len(result) == 5
        for img in result:
            assert img[:8] == b"\x89PNG\r\n\x1a\n"

    def test_max_pages(self) -> None:
        """A 15-page PDF with max_pages=10 returns only 10 images."""
        pdf_bytes = _make_pdf_pages(15)
        result = pdf_pages_to_images(pdf_bytes, max_pages=10)
        assert len(result) == 10

    def test_corrupted(self) -> None:
        """Invalid bytes raise ValueError."""
        with pytest.raises(ValueError, match="(?i)corrupt|cannot|open|invalid"):
            pdf_pages_to_images(b"this is not a pdf at all")

    def test_empty_pdf(self) -> None:
        """A 0-page PDF returns an empty list."""
        # Minimal valid PDF structure with zero pages (pymupdf cannot
        # serialize a 0-page document, so we build raw bytes).
        empty_pdf = (
            b"%PDF-1.0\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj "
            b"2 0 obj<</Type/Pages/Kids[]/Count 0>>endobj\n"
            b"xref\n0 3\n"
            b"0000000000 65535 f \n"
            b"0000000009 00000 n \n"
            b"0000000058 00000 n \n"
            b"trailer<</Size 3/Root 1 0 R>>\n"
            b"startxref\n109\n%%EOF"
        )
        result = pdf_pages_to_images(empty_pdf)
        assert result == []

    def test_dpi_parameter(self) -> None:
        """Higher DPI produces larger images than lower DPI."""
        pdf_bytes = _make_pdf("DPI test content")
        low_dpi = pdf_pages_to_images(pdf_bytes, dpi=72)
        high_dpi = pdf_pages_to_images(pdf_bytes, dpi=300)
        assert len(low_dpi) == 1
        assert len(high_dpi) == 1
        # Higher DPI should produce larger image data
        assert len(high_dpi[0]) > len(low_dpi[0])
